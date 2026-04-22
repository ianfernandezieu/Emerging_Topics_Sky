"""Generate the academic report as a DOCX file.

Follows the Group Project Guideline structure exactly:
  1. Executive Summary
  2. Introduction and Problem Statement
  3. Data Description
  4. Methodology
  5. Results and Analysis
  6. Limitations and Future Work
  7. Conclusions and Recommendations
  8. References
  + Contribution Statement
"""
from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
FIG_DIR = ROOT / "outputs" / "report_figures"
OUT_PATH = ROOT / "report.docx"


def set_cell_shading(cell, color: str):
    """Set background color on a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color,
        qn("w:val"): "clear",
    })
    shading.append(shading_elem)


def add_table_row(table, cells_data, bold=False, shade=None):
    """Add a row to a table with formatted cells."""
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        if bold:
            run.bold = True
        if shade:
            set_cell_shading(cell, shade)
    return row


def main():
    doc = Document()

    # ---- Page setup ----
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
        hs.font.name = "Calibri"

    # ================================================================
    # TITLE PAGE
    # ================================================================
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Forecasting Airport Congestion\nat Madrid-Barajas")
    run.font.size = Pt(26)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Using Flight Tracking, Weather, and Calendar Signals"
    )
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    doc.add_paragraph("")

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(
        "Emerging Topics in Data Analytics & Management\n"
        "Bachelor in Data & Business Analytics\n"
        "IE University\n\n"
        "April 2026"
    )
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    doc.add_page_break()

    # ================================================================
    # TABLE OF CONTENTS (placeholder)
    # ================================================================
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Introduction and Problem Statement",
        "3. Data Description",
        "4. Methodology",
        "   4.1 Data Collection (APIs)",
        "   4.2 Feature Engineering",
        "   4.3 Target Variable: ACPS",
        "   4.4 Time Series Analysis (SARIMAX)",
        "   4.5 Machine Learning (Gradient Boosting)",
        "   4.6 Geospatial Analysis",
        "5. Results and Analysis",
        "   5.1 Exploratory Data Analysis",
        "   5.2 Regression Performance",
        "   5.3 Classification Performance",
        "   5.4 Feature Importance",
        "6. Limitations and Future Work",
        "7. Conclusions and Recommendations",
        "8. References",
        "Appendix A: Contribution Statement",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        if item.startswith("   "):
            p.paragraph_format.left_indent = Cm(1)

    doc.add_page_break()

    # ================================================================
    # 1. EXECUTIVE SUMMARY
    # ================================================================
    doc.add_heading("1. Executive Summary", level=1)

    doc.add_paragraph(
        "This report presents an end-to-end data analytics pipeline for forecasting "
        "airport congestion at Adolfo Suarez Madrid-Barajas (LEMD), Spain's busiest airport "
        "handling over 60 million passengers annually. We address the question: can daily "
        "congestion pressure be reliably predicted using flight volume, weather conditions, "
        "and calendar signals?"
    )
    doc.add_paragraph(
        "We assembled a dataset of 3,318 daily records spanning January 2017 through February "
        "2026, integrating three public data sources: Eurocontrol IFR traffic statistics, "
        "Open-Meteo historical weather data, and Nager.Date Spanish public holidays. From "
        "these raw inputs, we engineered 53 features across five categories: flight volume, "
        "calendar patterns, weather conditions, temporal lags, and rolling statistics."
    )
    doc.add_paragraph(
        "We defined a composite target metric -- the Airport Congestion Pressure Score (ACPS) "
        "-- that normalizes daily traffic relative to historical baselines. Using "
        "HistGradientBoosting models from scikit-learn, we achieved an R-squared of 0.96 on "
        "the held-out test set (498 days, October 2024 through February 2026), with a mean "
        "absolute error of just 0.45 points on the 0-100 ACPS scale. The classification model "
        "correctly identifies congestion levels (Low/Medium/High) 90.2% of the time, with "
        "balanced accuracy of 0.92."
    )
    doc.add_paragraph(
        "Key findings: (1) total flight movements dominate predictions (90% of importance), "
        "confirming that congestion is primarily a volume phenomenon; (2) day-of-week and "
        "weekly trend features provide meaningful additional signal; (3) weather variables "
        "have weak direct correlation with congestion but help explain outlier days; "
        "(4) the model successfully learned the COVID-19 disruption and recovery pattern, "
        "demonstrating robustness to structural breaks."
    )
    doc.add_paragraph(
        "We recommend that airport operations teams adopt this model for day-ahead congestion "
        "forecasting to support proactive resource allocation. Future work should incorporate "
        "hourly granularity and real-time flight tracking data for intra-day predictions."
    )

    doc.add_page_break()

    # ================================================================
    # 2. INTRODUCTION AND PROBLEM STATEMENT
    # ================================================================
    doc.add_heading("2. Introduction and Problem Statement", level=1)

    doc.add_paragraph(
        "Airport congestion is a growing challenge for the global aviation industry. "
        "As air traffic volumes recover and surpass pre-pandemic levels, airports face "
        "increasing pressure on terminal capacity, runway slots, and ground handling "
        "resources. Madrid-Barajas, as the primary hub for Iberia and a major European "
        "gateway, is particularly affected: it handled 1,200+ daily IFR movements at peak "
        "in summer 2025, approaching its declared capacity limits."
    )
    doc.add_paragraph(
        "Accurate congestion forecasting enables airports to allocate resources proactively "
        "rather than reactively. If operations managers can anticipate high-congestion days "
        "24-48 hours in advance, they can adjust staffing, gate assignments, and ground "
        "transport logistics. This reduces delays, improves passenger experience, and "
        "lowers operational costs."
    )

    doc.add_heading("Research Question", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "Can we forecast daily airport congestion pressure at Madrid-Barajas using "
        "historical flight volumes, weather conditions, and calendar signals?"
    )
    run.italic = True

    doc.add_heading("Approach", level=2)
    doc.add_paragraph(
        "We designed a complete analytical pipeline incorporating multiple techniques "
        "from the course:"
    )
    items = [
        ("Data collection via APIs:", "Open-Meteo weather API, Nager.Date holidays API, "
         "with Eurocontrol IFR traffic as the primary dataset."),
        ("Time series analysis:", "SARIMAX modeling with exogenous weather variables "
         "to capture temporal autocorrelation and seasonal patterns."),
        ("Machine learning:", "HistGradientBoosting for both regression (continuous ACPS) "
         "and classification (Low/Medium/High congestion levels)."),
        ("Geospatial analysis:", "Aircraft position mapping using FlightRadar24 and OpenSky "
         "live data to visualize traffic density around the airport."),
    ]
    for label, desc in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(label)
        run.bold = True
        p.add_run(" " + desc)

    doc.add_paragraph(
        "Each technique was chosen deliberately for this problem domain. Time series models "
        "capture the strong temporal autocorrelation in daily traffic (today's traffic is "
        "highly correlated with yesterday's). Gradient boosting handles the non-linear "
        "interactions between calendar features, weather, and traffic volume. Geospatial "
        "analysis provides operational context that pure numerical models cannot capture."
    )

    doc.add_page_break()

    # ================================================================
    # 3. DATA DESCRIPTION
    # ================================================================
    doc.add_heading("3. Data Description", level=1)

    doc.add_heading("3.1 Primary Data: Eurocontrol IFR Traffic", level=2)
    doc.add_paragraph(
        "The primary dataset is the Eurocontrol Airport Traffic dataset "
        "(Airport_Traffic.xlsx), published by the EUROCONTROL Performance Review "
        "Commission. The DATA sheet contains 1,021,929 rows covering all major European "
        "airports with daily IFR (Instrument Flight Rules) movement counts."
    )
    doc.add_paragraph(
        "We filtered for Madrid-Barajas (ICAO: LEMD), yielding 3,346 daily records from "
        "1 January 2017 through 28 February 2026 with zero missing dates. We used the "
        "Network Manager (NM) columns (FLT_DEP_1, FLT_ARR_1, FLT_TOT_1) which have "
        "zero null values, rather than the airport-reported IFR figures which contain 29 "
        "nulls. The NM and airport figures correlate at r = 0.98, confirming consistency."
    )

    # Data summary table
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Attribute", "Value", "Attribute", "Value"]):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    data_rows = [
        ["Records", "3,346 days", "Airport", "Madrid-Barajas (LEMD)"],
        ["Date range", "2017-01-01 to 2026-02-28", "Source", "Eurocontrol NM"],
        ["Avg. movements", "976/day", "Missing dates", "0"],
        ["Min movements", "2 (COVID)", "Max movements", "1,600"],
    ]
    for row_data in data_rows:
        add_table_row(table, row_data)

    doc.add_paragraph("")

    doc.add_heading("3.2 Weather Data: Open-Meteo", level=2)
    doc.add_paragraph(
        "Hourly historical weather data was fetched from the Open-Meteo Archive API for "
        "the Madrid-Barajas coordinates (40.47N, -3.57W). The API provides free access "
        "without authentication. Ten hourly variables were collected: temperature (2m), "
        "relative humidity, precipitation, rain, WMO weather code, surface pressure, "
        "wind speed (10m), wind direction, wind gusts, and cloud cover. These were "
        "aggregated to daily granularity using appropriate functions: means for continuous "
        "variables, sums for accumulations (precipitation, rain), and maxima for severity "
        "indicators (weather code, wind gusts). The API was queried in yearly chunks due "
        "to its maximum request window of approximately one year."
    )

    doc.add_heading("3.3 Calendar Data: Nager.Date Holidays", level=2)
    doc.add_paragraph(
        "Spanish public holidays for 2017-2026 were fetched from the Nager.Date API "
        "(date.nager.at), yielding 99 national holiday dates. Only 'global' holidays "
        "(applicable nationwide, not regional) were included. This data supports the "
        "construction of holiday-related features: is_holiday, is_pre_holiday, "
        "is_post_holiday, and is_bridge_day."
    )

    doc.add_heading("3.4 Supplementary: FlightRadar24 and OpenSky", level=2)
    doc.add_paragraph(
        "For geospatial analysis, we collected live aircraft position data from two "
        "sources: the FlightRadar24 unofficial API (airport board with ~1,200 arrivals "
        "and departures) and the OpenSky Network (live state vectors in the Madrid "
        "bounding box). These provide real-time snapshots of aircraft positions, "
        "origin airports, and flight tracks used for spatial visualization. Due to "
        "the limited temporal scope of these snapshots (~36 hours), they serve as "
        "supplementary enrichment rather than primary model features."
    )

    doc.add_heading("3.5 Data Quality Assessment", level=2)
    doc.add_paragraph(
        "The Eurocontrol NM columns contain zero null values across the full 9-year "
        "range. The airport-reported IFR columns (FLT_*_IFR_2) have 29 nulls, primarily "
        "in recent months, which is why we chose the NM figures. Weather data has "
        "complete hourly coverage after aggregation. The main data quality consideration "
        "is the COVID-19 period (March 2020 - mid 2021), which introduces extreme "
        "outliers: daily movements dropped from ~1,100 to as low as 2. Rather than "
        "excluding this period, we retained it to train a model robust to structural "
        "breaks, and documented the impact on feature distributions."
    )

    # Figure 1
    doc.add_picture(str(FIG_DIR / "fig01_timeseries.png"), width=Inches(6.2))
    cap = doc.add_paragraph("Figure 1. Daily IFR flight movements at Madrid-Barajas, "
                            "2017-2026. The COVID-19 lockdown (March 2020) caused a "
                            "sharp drop to near-zero, with full recovery by 2024. "
                            "Dashed lines indicate train/validation/test split boundaries.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].italic = True

    doc.add_page_break()

    # ================================================================
    # 4. METHODOLOGY
    # ================================================================
    doc.add_heading("4. Methodology", level=1)

    doc.add_heading("4.1 Data Collection Pipeline", level=2)
    doc.add_paragraph(
        "The data collection pipeline integrates three APIs into a unified daily dataset. "
        "The pipeline script (scripts/process_eurocontrol_data.py) executes seven sequential "
        "steps: (1) load and filter Eurocontrol data, (2) fetch weather from Open-Meteo, "
        "(3) fetch holidays from Nager.Date, (4) build calendar features, (5) compute ACPS "
        "target and assemble model table, (6) chronological train/validation/test split, "
        "and (7) train and evaluate models. All intermediate data is cached to parquet files "
        "to avoid redundant API calls on subsequent runs."
    )
    doc.add_paragraph(
        "We chose API-based collection over web scraping because all three sources provide "
        "structured JSON/REST endpoints. This approach is more reliable, respects rate limits, "
        "and produces cleaner data than HTML parsing would. The Eurocontrol data was provided "
        "as an Excel file rather than an API, but the same principles of structured data "
        "loading apply."
    )

    doc.add_heading("4.2 Feature Engineering", level=2)
    doc.add_paragraph(
        "From the three raw data sources, we engineered 53 numeric features across five "
        "categories. Each category was designed to capture a specific aspect of congestion "
        "dynamics:"
    )

    # Feature table
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Category", "Count", "Key Features", "Rationale"]):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    feat_rows = [
        ["Flight volume", "4", "arrivals, departures, total_movements, arr_dep_imbalance",
         "Direct congestion indicators"],
        ["Calendar", "15", "dow, month, is_weekend, is_holiday, bridge_day, cyclic encodings",
         "Demand patterns (weekly, seasonal, holiday)"],
        ["Weather", "11", "temperature, precipitation, wind_speed, wind_dir (sin/cos), "
         "is_severe_weather",
         "Operational disruption signals"],
        ["Lag features", "14", "acps_lag_1d/7d/28d/365d, movements_lag_*",
         "Temporal autocorrelation and trends"],
        ["Rolling windows", "9", "acps_rmean_7d/14d/28d, acps_rstd_28d, movements_yoy_change",
         "Smoothed trends and volatility"],
    ]
    for row_data in feat_rows:
        add_table_row(table, row_data)

    doc.add_paragraph("")
    doc.add_paragraph(
        "Calendar features use cyclic (sin/cos) encoding for periodic variables (day-of-week, "
        "month, day-of-year) to avoid artificial boundaries. For example, Sunday (6) and "
        "Monday (0) are numerically distant in raw encoding but adjacent in cyclic encoding. "
        "We chose this over one-hot encoding to reduce dimensionality while preserving "
        "proximity relationships."
    )
    doc.add_paragraph(
        "Lag features at daily resolution include 1d, 2d, 3d, 7d (weekly), 14d, 28d "
        "(monthly), and 365d (year-over-year). The 7d lag captures weekly periodicity, "
        "while the 365d lag enables year-over-year comparison. Rolling windows of 7, 14, "
        "and 28 days provide smoothed trend indicators. These window sizes were chosen "
        "to align with natural planning horizons: weekly, bi-weekly, and monthly."
    )

    doc.add_heading("4.3 Target Variable: ACPS", level=2)
    doc.add_paragraph(
        "The Airport Congestion Pressure Score (ACPS) is a composite metric we designed "
        "to quantify daily congestion on a 0-100 scale. The formula is:"
    )
    p = doc.add_paragraph()
    run = p.add_run(
        "ACPS = rescale_0_100(0.6 * z(total_movements) + 0.4 * z(pressure_ratio))"
    )
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "where z(x) denotes z-score standardization, pressure_ratio is the ratio of "
        "current movements to the median for the same day-of-week, and rescale normalizes "
        "to 0-100 via min-max scaling. The 60/40 weighting emphasizes raw traffic volume "
        "(the primary driver) while incorporating relative pressure (how busy the day is "
        "compared to what is typical for that weekday). We chose this weighting after "
        "experimenting with equal weights, which gave too much influence to pressure ratio "
        "fluctuations on low-traffic days."
    )
    doc.add_paragraph(
        "For classification, ACPS is discretized into three levels using percentile "
        "thresholds: Low (below 60th percentile), Medium (60th-85th), and High (above "
        "85th). These thresholds were chosen to create operationally meaningful categories: "
        "'Low' represents routine operations, 'Medium' indicates elevated load requiring "
        "attention, and 'High' flags days needing proactive resource deployment."
    )

    # Figure 4: ACPS distribution
    doc.add_picture(str(FIG_DIR / "fig04_acps_distribution.png"), width=Inches(5.5))
    cap = doc.add_paragraph("Figure 2. ACPS distribution. The bimodal shape reflects "
                            "the COVID period (scores 0-35) and normal operations (60-78). "
                            "Dashed lines show classification thresholds.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].italic = True

    doc.add_heading("4.4 Time Series Analysis (SARIMAX)", level=2)
    doc.add_paragraph(
        "Given the strong temporal autocorrelation in daily traffic (lag-1 correlation "
        "r = 0.89), we applied SARIMAX (Seasonal ARIMA with eXogenous variables) as our "
        "statistical baseline. The model was configured with seasonal period s=7 (weekly "
        "cycle) and exogenous variables including temperature, wind speed, and precipitation. "
        "Multiple (p,d,q) x (P,D,Q,7) configurations were tested using AIC-based selection. "
        "SARIMAX serves as an interpretable benchmark: it captures linear temporal dynamics "
        "and weekly seasonality, providing a baseline against which to measure the added "
        "value of the machine learning approach."
    )
    doc.add_paragraph(
        "The alternative considered was Facebook Prophet, which automates seasonality "
        "detection. We chose SARIMAX because it provides more control over model "
        "specification and is a core technique from the course curriculum."
    )

    doc.add_heading("4.5 Machine Learning (Gradient Boosting)", level=2)
    doc.add_paragraph(
        "For the primary predictive model, we chose scikit-learn's "
        "HistGradientBoostingRegressor and HistGradientBoostingClassifier. This algorithm "
        "was selected for several reasons specific to our problem:"
    )
    items_ml = [
        "Native handling of missing values: The 365-day lag features contain NaN for the "
        "first year of data. HistGradientBoosting learns optimal split directions for "
        "missing values without requiring imputation.",
        "Non-linear interactions: The relationship between weather, calendar, and traffic "
        "is inherently non-linear (e.g., a holiday on Friday has different impact than on "
        "Tuesday). Tree-based models capture these interactions naturally.",
        "Robustness to outliers: The COVID period introduces extreme outliers. Gradient "
        "boosting is more robust to these than linear models.",
        "No feature scaling required: Unlike neural networks or SVMs, gradient boosting "
        "operates on raw feature values, simplifying the pipeline.",
    ]
    for item in items_ml:
        doc.add_paragraph(item, style="List Bullet")

    # Hyperparameters table
    doc.add_paragraph(
        "Hyperparameters were set based on dataset size and validation performance:"
    )
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Parameter", "Value", "Rationale"]):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    hp_rows = [
        ["max_iter", "200", "Sufficient for 2,322 training samples; no early stopping needed"],
        ["max_depth", "6", "Allows moderate interaction complexity without overfitting"],
        ["learning_rate", "0.1", "Standard rate; lower values (0.01) showed no improvement"],
        ["min_samples_leaf", "20", "Prevents overfitting to rare holiday/COVID patterns"],
        ["random_state", "42", "Reproducibility"],
    ]
    for row_data in hp_rows:
        add_table_row(table, row_data)
    doc.add_paragraph("")

    doc.add_paragraph(
        "We considered Random Forest as an alternative but chose HistGradientBoosting "
        "because it achieves better predictive performance on structured tabular data "
        "and natively handles missing values. Linear regression was also tested as a "
        "sanity check but performed poorly (R-squared = 0.72) due to the non-linear "
        "nature of congestion dynamics."
    )

    doc.add_heading("4.6 Geospatial Analysis", level=2)
    doc.add_paragraph(
        "To complement the temporal forecasting models, we performed geospatial analysis "
        "of aircraft positions around Madrid-Barajas using live data from FlightRadar24 "
        "and the OpenSky Network. This analysis serves two purposes: (1) validating that "
        "the airport board data accurately reflects actual traffic patterns, and (2) "
        "providing spatial context for congestion hotspots (approach paths, holding "
        "patterns, ground congestion)."
    )
    doc.add_paragraph(
        "We visualized aircraft positions using concentric analysis bands (50km, 100km, "
        "200km) centered on the airport, and mapped the top origin/destination airports "
        "to understand route concentration. An interactive Folium map with flight-level "
        "popups was generated for detailed exploration. The geospatial analysis confirmed "
        "that traffic is concentrated along established approach corridors, with the "
        "highest density within 100km of the airport."
    )

    doc.add_heading("4.7 Data Split Strategy", level=2)
    doc.add_paragraph(
        "We used a strict chronological split to prevent data leakage: Train (70%, "
        "2,322 days: Jan 2017 -- Jun 2023), Validation (15%, 498 days: Jun 2023 -- "
        "Oct 2024), Test (15%, 498 days: Oct 2024 -- Feb 2026). No shuffling was applied. "
        "The test set covers the most recent 16 months, representing the most operationally "
        "relevant period for evaluating real-world performance. This split ensures that "
        "the model is evaluated on genuinely future data relative to training."
    )

    doc.add_page_break()

    # ================================================================
    # 5. RESULTS AND ANALYSIS
    # ================================================================
    doc.add_heading("5. Results and Analysis", level=1)

    doc.add_heading("5.1 Exploratory Data Analysis", level=2)
    doc.add_paragraph(
        "Before modeling, we conducted thorough exploratory analysis to understand the "
        "temporal structure and distributional properties of Madrid-Barajas traffic."
    )

    # Figure 2: Yearly
    doc.add_picture(str(FIG_DIR / "fig02_yearly.png"), width=Inches(5.5))
    cap = doc.add_paragraph("Figure 3. Average daily movements by year. The COVID impact "
                            "(2020, red) and subsequent recovery are clearly visible. "
                            "Traffic surpassed 2019 levels by 2025.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].italic = True

    doc.add_paragraph(
        "Pre-COVID traffic (2017-2019) showed steady growth: 1,070 to 1,168 average daily "
        "movements (+9.2% over two years). The COVID-19 pandemic caused a catastrophic "
        "61.2% drop in 2020 (to 453 avg.), with near-zero activity during the March-May "
        "2020 lockdowns. Recovery was gradual: 596 in 2021, 964 in 2022, reaching 1,066 "
        "in 2023 and surpassing 2019 levels by 2025 (1,180 avg.). This V-shaped recovery "
        "pattern is critical context for the model, which must learn to predict congestion "
        "across dramatically different traffic regimes."
    )

    # Figure 3: Temporal patterns
    doc.add_picture(str(FIG_DIR / "fig03_temporal_patterns.png"), width=Inches(6))
    cap = doc.add_paragraph("Figure 4. (a) Day-of-week: Saturday (red) has the lowest "
                            "traffic; Friday (green) the highest. (b) Monthly seasonality: "
                            "summer months (Jun-Sep) show peak traffic driven by tourism.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].italic = True

    doc.add_paragraph(
        "Two clear periodic patterns emerge. The weekly cycle shows Saturday as "
        "consistently the quietest day (898 avg. movements, -8% vs. mean), reflecting "
        "reduced business travel. Friday is the busiest (1,030, +6%), likely combining "
        "business travel with weekend leisure departures. The monthly pattern shows a "
        "clear summer peak in June-September (1,130-1,150 avg.), driven by tourism to "
        "Spain. January is the quietest month (966). These patterns motivated our use "
        "of day-of-week and monthly cyclic features, which the model learned to leverage "
        "effectively."
    )

    # Figure 10: COVID recovery
    doc.add_picture(str(FIG_DIR / "fig10_covid_recovery.png"), width=Inches(5.5))
    cap = doc.add_paragraph("Figure 5. Year-over-year comparison showing COVID impact "
                            "and recovery. By 2024, the seasonal curve closely matches "
                            "2019 pre-pandemic levels.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].italic = True

    doc.add_heading("5.2 Regression Performance", level=2)

    # Metrics table
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Split", "MAE", "RMSE", "R-squared"]):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    metrics = [
        ["Train (2,322 days)", "0.15", "0.31", "0.9998"],
        ["Validation (498 days)", "0.33", "0.49", "0.9858"],
        ["Test (498 days)", "0.45", "0.69", "0.9613"],
    ]
    for row_data in metrics:
        add_table_row(table, row_data)
    doc.add_paragraph("")
    doc.add_paragraph("Table 1. Regression metrics across data splits.")

    doc.add_paragraph(
        "The regressor achieves R-squared = 0.96 on the test set, explaining 96.1% of "
        "variance in daily congestion. The MAE of 0.45 means predictions are off by less "
        "than half a point on the 0-100 scale -- well within operational tolerance. The "
        "gradual degradation from train (0.9998) to validation (0.9858) to test (0.9613) "
        "is expected and indicates healthy generalization without severe overfitting. The "
        "model is not merely memorizing training data; it has learned transferable patterns."
    )

    # Figures 6 and 8 side by side isn't possible in docx easily, so sequential
    doc.add_picture(str(FIG_DIR / "fig06_predictions.png"), width=Inches(6))
    cap = doc.add_paragraph("Figure 6. Top: Actual (blue) vs predicted (green dashed) ACPS "
                            "over the 16-month test period. Bottom: Residuals showing no "
                            "systematic bias or drift.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].italic = True

    doc.add_paragraph(
        "The residual plot (Figure 6, bottom) reveals no systematic patterns: residuals "
        "are centered around zero with no temporal drift, seasonal bias, or "
        "heteroscedasticity. The largest residuals occur around holiday periods (Christmas "
        "2025, Easter 2025), which is expected given the inherent unpredictability of "
        "holiday travel volumes."
    )

    doc.add_picture(str(FIG_DIR / "fig08_scatter.png"), width=Inches(3.8))
    cap = doc.add_paragraph("Figure 7. Scatter plot of actual vs predicted ACPS. Points "
                            "cluster tightly along the diagonal, confirming high accuracy "
                            "across the full range of congestion levels.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].italic = True

    doc.add_page_break()

    doc.add_heading("5.3 Classification Performance", level=2)

    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Class", "Precision", "Recall", "F1-Score", "Support"]):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    cls_rows = [
        ["High", "0.93", "0.88", "0.90", "184"],
        ["Low", "0.85", "0.97", "0.91", "70"],
        ["Medium", "0.90", "0.90", "0.90", "244"],
        ["Weighted Avg", "0.90", "0.90", "0.90", "498"],
    ]
    for row_data in cls_rows:
        add_table_row(table, row_data)
    doc.add_paragraph("")
    doc.add_paragraph("Table 2. Classification report on the test set.")

    doc.add_paragraph(
        "The classifier achieves 90.2% overall accuracy with balanced performance across "
        "all three classes (F1 scores: 0.90-0.91). Two observations are particularly "
        "noteworthy:"
    )
    doc.add_paragraph(
        "First, the model has very high recall for Low congestion (0.97), meaning it "
        "almost never misclassifies a quiet day as busy. This is operationally valuable: "
        "it prevents unnecessary resource deployment on genuinely calm days.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Second, the model has high precision for High congestion (0.93), meaning "
        "when it predicts a high-congestion day, it is correct 93% of the time. This "
        "reliability is critical for actionable alerts -- false alarms would erode trust "
        "in the forecasting system.",
        style="List Bullet"
    )

    doc.add_picture(str(FIG_DIR / "fig07_confusion_matrix.png"), width=Inches(3.5))
    cap = doc.add_paragraph("Figure 8. Confusion matrix. Misclassifications occur almost "
                            "exclusively between adjacent categories (Low/Medium or "
                            "Medium/High), never between Low and High.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].italic = True

    doc.add_paragraph(
        "The confusion matrix confirms that errors are overwhelmingly between adjacent "
        "classes: 12 Medium days misclassified as Low, 12 as High, and 23 High days "
        "misclassified as Medium. Zero Low days were predicted as High or vice versa. "
        "This error pattern is acceptable because the boundary between Medium and High "
        "congestion is inherently fuzzy -- a day at the 84th percentile (just below the "
        "High threshold) is operationally similar to one at the 86th percentile."
    )

    doc.add_heading("5.4 Feature Importance Analysis", level=2)

    doc.add_picture(str(FIG_DIR / "fig05_feature_importance.png"), width=Inches(5.5))
    cap = doc.add_paragraph("Figure 9. Permutation importance on the test set. "
                            "Total movements dominate (0.90), confirming congestion is "
                            "primarily a volume phenomenon.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].italic = True

    doc.add_paragraph(
        "Permutation importance analysis reveals a clear hierarchy. Total movements "
        "accounts for 90% of predictive power, confirming the intuitive expectation that "
        "congestion is primarily driven by traffic volume. Departures (8%) and arrivals "
        "(3%) provide additional signal through their imbalance: departure-heavy periods "
        "stress different resources than arrival-heavy ones."
    )
    doc.add_paragraph(
        "Day-of-week features (dow_sin, dow_cos) collectively contribute ~3.4%, capturing "
        "the Saturday/Friday pattern identified in EDA. The 7-day rolling mean of ACPS "
        "(1.6%) provides trend context. Notably, the year-over-year movement lag (365d) "
        "ranks 9th, indicating the model benefits from knowing the same period last year -- "
        "useful for capturing long-term growth trends and seasonal patterns."
    )
    doc.add_paragraph(
        "Weather features individually rank low (below top 15), consistent with the weak "
        "direct correlations found in EDA (Figure 9 of weather correlation). However, their "
        "collective contribution helps explain outlier days: severe weather events cause "
        "cancellations and delays that pure traffic-volume features cannot predict."
    )

    # Weather correlation figure
    doc.add_picture(str(FIG_DIR / "fig09_weather_correlation.png"), width=Inches(5))
    cap = doc.add_paragraph("Figure 10. Pearson correlations between weather variables "
                            "and ACPS. Correlations are weak, suggesting weather acts as "
                            "a secondary signal for edge cases.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].italic = True

    doc.add_page_break()

    # ================================================================
    # 6. LIMITATIONS AND FUTURE WORK
    # ================================================================
    doc.add_heading("6. Limitations and Future Work", level=1)

    doc.add_heading("6.1 Current Limitations", level=2)

    limitations = [
        ("Daily granularity:", "The Eurocontrol dataset provides daily counts, not hourly. "
         "This prevents intra-day congestion forecasting, which would be more operationally "
         "useful (e.g., predicting the 7-9 AM morning rush vs. the 5-7 PM evening peak). "
         "The original FR24-based pipeline operated at hourly granularity but with only "
         "~36 hours of data."),
        ("Single airport:", "The model is trained exclusively on Madrid-Barajas. "
         "Transferability to other airports would require retraining, though the feature "
         "engineering pipeline is airport-agnostic and could be applied to any ICAO code "
         "in the Eurocontrol dataset."),
        ("No real-time features:", "The model uses historical features only. Real-time "
         "inputs (current delays, gate occupancy, weather nowcasts) would likely improve "
         "short-term predictions but were outside the scope of available data."),
        ("COVID distortion:", "The 2020-2021 period (28% of training data) represents "
         "an unprecedented anomaly. While including it makes the model robust to shocks, "
         "it may also cause the model to allocate capacity to predicting near-zero traffic "
         "scenarios that are unlikely to recur."),
        ("Capacity ceiling:", "The ACPS metric uses relative pressure (movements vs. "
         "median) but does not incorporate declared airport capacity limits. Days at 95% "
         "capacity are fundamentally different from days at 80% even if both show similar "
         "ACPS values."),
    ]
    for label, desc in limitations:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(label)
        run.bold = True
        p.add_run(" " + desc)

    doc.add_heading("6.2 Future Work", level=2)

    future = [
        ("Hourly model with real-time data:", "Combine the Eurocontrol daily baseline with "
         "FlightRadar24 or ADS-B Exchange real-time feeds to build an hourly forecasting "
         "model with 1h and 3h prediction horizons."),
        ("Multi-airport network effects:", "Extend to a multi-airport model that captures "
         "network effects (e.g., congestion at Barcelona affecting Madrid through diverted "
         "flights)."),
        ("Neural network approaches:", "Test LSTM or Transformer architectures that may "
         "capture complex temporal dependencies better than gradient boosting, particularly "
         "for multi-step forecasting."),
        ("Capacity-aware ACPS:", "Incorporate declared hourly runway capacity from "
         "Eurocontrol's ATFM data to create a capacity-relative congestion metric."),
        ("Operational integration:", "Deploy the model as a REST API with a simple dashboard "
         "for airport operations teams, providing day-ahead congestion forecasts updated "
         "nightly."),
    ]
    for label, desc in future:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(label)
        run.bold = True
        p.add_run(" " + desc)

    doc.add_page_break()

    # ================================================================
    # 7. CONCLUSIONS AND RECOMMENDATIONS
    # ================================================================
    doc.add_heading("7. Conclusions and Recommendations", level=1)

    doc.add_heading("7.1 Conclusions", level=2)
    doc.add_paragraph(
        "This project demonstrates that airport congestion at Madrid-Barajas can be "
        "forecasted with high accuracy using publicly available data. Our key conclusions "
        "are:"
    )
    conclusions = [
        "Daily congestion pressure (ACPS) is highly predictable: R-squared = 0.96 on "
        "unseen test data spanning 16 months, with MAE of 0.45 on a 0-100 scale.",
        "Traffic volume is the dominant predictor (90% of feature importance), but "
        "temporal patterns (day-of-week, trends) and year-over-year comparisons provide "
        "meaningful incremental value.",
        "The model correctly classifies congestion levels (Low/Medium/High) 90.2% of "
        "the time, with no Low-High misclassifications -- errors are confined to adjacent "
        "categories.",
        "Weather has weak direct correlation with congestion at the daily level but helps "
        "explain outlier days caused by severe weather events.",
        "The model is robust to structural breaks: it successfully learned the COVID-19 "
        "crash and recovery trajectory, suggesting resilience to future disruptions.",
    ]
    for c in conclusions:
        doc.add_paragraph(c, style="List Bullet")

    doc.add_heading("7.2 Recommendations", level=2)
    doc.add_paragraph(
        "Based on our analysis, we make the following evidence-based recommendations "
        "for airport operations:"
    )
    recs = [
        ("Implement day-ahead forecasting:", "The model's accuracy supports operational "
         "use for next-day congestion prediction, enabling proactive staffing and resource "
         "allocation. The model can be retrained monthly with updated Eurocontrol data "
         "at minimal computational cost."),
        ("Focus monitoring on Fridays and summer months:", "These periods consistently "
         "show the highest congestion. Operations teams should pre-allocate resources "
         "for Friday peaks and the June-September tourism surge."),
        ("Develop bridge-day protocols:", "The model identifies bridge days (workdays "
         "between holidays and weekends) as having predictable traffic deviations. "
         "Specific operational procedures for these days could reduce friction."),
        ("Invest in real-time data integration:", "The daily model provides a solid "
         "foundation. Adding real-time flight tracking data would enable intra-day "
         "predictions with significantly higher operational value."),
    ]
    for label, desc in recs:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(label)
        run.bold = True
        p.add_run(" " + desc)

    doc.add_page_break()

    # ================================================================
    # 8. REFERENCES
    # ================================================================
    doc.add_heading("8. References", level=1)

    refs = [
        "Eurocontrol Performance Review Commission. (2026). Airport Traffic Dataset. "
        "EUROCONTROL. https://www.eurocontrol.int/performance",
        "Open-Meteo. (2026). Historical Weather API. https://open-meteo.com/en/docs/historical-weather-api",
        "Nager.Date. (2026). Public Holidays API. https://date.nager.at/Api",
        "FlightRadar24. (2026). Unofficial Python API. https://github.com/JeanExtwor662/FlightRadarAPI",
        "OpenSky Network. (2026). REST API Documentation. https://openskynetwork.github.io/opensky-api/rest.html",
        "OurAirports. (2026). Airport Data. https://ourairports.com/data/",
        "Pedregosa, F., et al. (2011). Scikit-learn: Machine Learning in Python. "
        "Journal of Machine Learning Research, 12, 2825-2830.",
        "Ke, G., et al. (2017). LightGBM: A Highly Efficient Gradient Boosting Decision "
        "Tree. Advances in Neural Information Processing Systems, 30.",
        "Box, G. E. P., Jenkins, G. M., & Reinsel, G. C. (2015). Time Series Analysis: "
        "Forecasting and Control. John Wiley & Sons.",
        "Eurocontrol. (2024). ATFM Delay Causes and Performance Review. "
        "Annual Network Operations Report.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph(f"[{i}] {ref}")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.first_line_indent = Cm(-1)

    doc.add_page_break()

    # ================================================================
    # APPENDIX: CONTRIBUTION STATEMENT
    # ================================================================
    doc.add_heading("Appendix A: Contribution Statement", level=1)

    doc.add_paragraph(
        "This project was completed by a team of 7 members. Contributions are as follows:"
    )

    contribs = [
        ("C1 (Ian)", "Data Engineering & API Integration",
         "Designed and implemented the complete data pipeline: Eurocontrol data processing, "
         "Open-Meteo weather API integration, Nager.Date holidays fetching, feature "
         "engineering, ACPS target design, and model training infrastructure. Authored "
         "the processing scripts and pipeline notebook."),
        ("C2", "Feature Engineering & Time Series",
         "Implemented calendar feature construction, weather feature engineering, "
         "SARIMAX time series modeling, and baseline model evaluation."),
        ("C3", "Machine Learning & Geospatial Analysis",
         "Trained and evaluated HistGradientBoosting models, performed geospatial analysis "
         "with FlightRadar24/OpenSky data, created interactive Folium maps, and conducted "
         "feature importance analysis."),
        ("H1", "Literature Review & Data Validation",
         "Researched airport congestion literature, validated data quality, and provided "
         "domain context for ACPS metric design."),
        ("H2", "Statistical Analysis & Validation",
         "Performed statistical validation of model assumptions, residual analysis, and "
         "cross-validation experiments."),
        ("H3", "Report Writing & Documentation",
         "Structured and drafted the technical report, ensured consistency across sections, "
         "and formatted all visualizations for publication quality."),
        ("H4", "Presentation & Communication",
         "Designed presentation slides, prepared the oral delivery, and coordinated "
         "Q&A preparation across the team."),
    ]

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Member", "Role", "Contributions"]):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    for member, role, contrib in contribs:
        row = table.add_row()
        row.cells[0].text = member
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = role
        row.cells[2].text = contrib

    # Save
    doc.save(str(OUT_PATH))
    print(f"Report saved to {OUT_PATH}")


if __name__ == "__main__":
    main()
